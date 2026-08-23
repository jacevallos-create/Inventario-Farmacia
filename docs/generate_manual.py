from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("Manual_Funcional_PharmaSys_Actualizado.docx")
NAVY = "173B57"
BLUE = "2878B5"
TEAL = "25A6A1"
LIGHT = "EAF3F7"
PALE = "F4F7F9"
GRAY = "667784"
WHITE = "FFFFFF"
INK = "20303B"
WIDTHS = {2: [2700, 6660], 3: [2050, 3655, 3655], 4: [1800, 2380, 2380, 2800]}


def font(run, name="Aptos", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if shd.getparent() is None:
        tc_pr.append(shd)


def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}")) or OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        if node.getparent() is None:
            tc_mar.append(node)


def table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)


def set_cell_text(cell, text, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(str(text))
    font(r, size=9.2, color=color, bold=bold)


def add_table(doc, headers, rows, widths=None):
    widths = widths or WIDTHS[len(headers)]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table_geometry(table, widths)
    for idx, value in enumerate(headers):
        shade(table.rows[0].cells[idx], NAVY)
        set_cell_text(table.rows[0].cells[idx], value, bold=True, color=WHITE)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            if row_index % 2:
                shade(cells[idx], PALE)
            set_cell_text(cells[idx], value)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    font(r, size=10.5, color=INK)
    return p


def add_step(doc, title, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_together = True
    r = p.add_run(f"{title}. ")
    font(r, size=10.5, color=NAVY, bold=True)
    r = p.add_run(text)
    font(r, size=10.5, color=INK)


def callout(doc, label, text, fill=LIGHT):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.14)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), TEAL)
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(f"{label}: ")
    font(r, size=10.2, color=NAVY, bold=True)
    r = p.add_run(text)
    font(r, size=10.2, color=INK)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.keep_together = True
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.bottom_margin = Inches(0.82)
section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.42)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15
for name, size, before, after, color in (
    ("Heading 1", 16, 18, 8, NAVY),
    ("Heading 2", 13, 13, 6, BLUE),
    ("Heading 3", 11.5, 9, 4, NAVY),
):
    style = styles[name]
    style.font.name = "Aptos Display"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
for name in ("List Bullet", "List Bullet 2", "List Number"):
    style = styles[name]
    style.font.name = "Aptos"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

header = section.header
hp = header.paragraphs[0]
hp.text = "PHARMASYS  |  MANUAL FUNCIONAL"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(hp.runs[0], size=8.5, color=GRAY, bold=True)
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(fp.add_run("PharmaSys | Documento operativo | "), size=8.5, color=GRAY)
field = OxmlElement("w:fldSimple")
field.set(qn("w:instr"), "PAGE")
fp._p.append(field)

# Cover
for _ in range(5):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PHARMASYS")
font(r, size=13, color=TEAL, bold=True)
p.paragraph_format.space_after = Pt(14)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Manual funcional del sistema")
font(r, name="Aptos Display", size=30, color=NAVY, bold=True)
p.paragraph_format.space_after = Pt(8)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Inventario, compras, ventas, caja y control multi-sucursal")
font(r, size=14, color=BLUE)
p.paragraph_format.space_after = Pt(32)
callout(doc, "Objetivo", "Explicar las funciones disponibles, los permisos por rol y el flujo correcto de operación en PharmaSys.")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(44)
r = p.add_run("Versión funcional - agosto de 2026")
font(r, size=10, color=GRAY, italic=True)
doc.add_page_break()

heading(doc, "1. Visión general", 1)
body(doc, "PharmaSys centraliza la operación de una o varias farmacias. La información persistente se almacena en PostgreSQL de Supabase y la aplicación se ejecuta en Django con una interfaz web adaptable a computadoras, tabletas y teléfonos.")
callout(doc, "Regla de acceso", "La pantalla inicial permite ingresar con correo y contraseña o continuar con Google. Las credenciales de demostración no se muestran públicamente. El administrador registra cada colaborador y asigna su rol y farmacias.")
heading(doc, "1.1 Flujo general de uso", 2)
add_step(doc, "Acceder", "Presione Continuar con Google y seleccione la cuenta autorizada.")
add_step(doc, "Seleccionar farmacia", "Use el selector superior para trabajar en una sucursal asignada.")
add_step(doc, "Operar", "Abra los módulos disponibles según su rol y registre cada movimiento con datos reales.")
add_step(doc, "Confirmar", "Las acciones sensibles usan cuadros de diálogo del sistema con motivo o autorización cuando corresponde.")
add_step(doc, "Cerrar sesión", "Abra el perfil y confirme la salida desde el diálogo visual de PharmaSys.")

heading(doc, "2. Roles y permisos", 1)
body(doc, "Los permisos se calculan al iniciar sesión a partir de la asignación activa del usuario. Un colaborador solo ve las farmacias y módulos que le corresponden.")
add_table(doc, ["Rol", "Módulos visibles", "Uso principal"], [
    ("Administrador", "Todos", "Configuración, usuarios, sucursales, operación y autorizaciones."),
    ("Farmacéutico", "Dashboard, Inventario, POS, Proveedores, Compras, Transferencias y Caja", "Atención, abastecimiento y operación diaria."),
    ("Inventario", "Dashboard, Inventario, Proveedores, Compras y Transferencias", "Existencias, recepción, lotes y movimientos entre sedes."),
    ("Cajero", "Dashboard, POS y Caja", "Ventas, medios de pago y control de turno."),
    ("Consulta", "Dashboard e Inventario", "Lectura de indicadores y existencias sin administración."),
], [1900, 4160, 3300])
callout(doc, "Autorizaciones", "Anular ventas, anular órdenes, devolver a proveedores, completar acciones administrativas de transferencias, administrar usuarios y exportar ciertos reportes requiere un administrador.", "FFF4D8")

heading(doc, "3. Inicio de sesión y seguridad", 1)
heading(doc, "3.1 Acceso al sistema", 2)
add_bullet(doc, "Una cuenta creada por el administrador puede ingresar con su correo y contraseña temporal.")
add_bullet(doc, "Google verifica la identidad y devuelve al usuario a PharmaSys.")
add_bullet(doc, "Para usar Google, el correo debe coincidir exactamente con el registrado por el administrador.")
add_bullet(doc, "Si cambia el rol o las sucursales, cierre sesión y vuelva a entrar para actualizar los permisos.")
heading(doc, "3.2 Gestión de usuarios", 2)
add_step(doc, "Crear", "Ingrese nombre y correo, seleccione un rol y asigne una o varias farmacias.")
add_step(doc, "Editar", "Cambie el rol, estado o las sucursales autorizadas desde el botón de edición.")
add_step(doc, "Eliminar", "Confirme la eliminación en el cuadro visual. El sistema protege al último administrador activo.")
callout(doc, "Buena práctica", "Asigne el rol de menor privilegio que permita realizar el trabajo. No comparta cuentas entre colaboradores.")

heading(doc, "4. Dashboard", 1)
body(doc, "El panel utiliza información real de Supabase para resumir ventas, utilidad, inventario valorizado, pérdidas por vencimiento y comparación entre farmacias.")
add_bullet(doc, "Indicadores de ventas e inventario para la farmacia seleccionada.")
add_bullet(doc, "Alertas rápidas de stock mínimo y lotes próximos a vencer.")
add_bullet(doc, "Comparación operativa entre sucursales para administradores.")
add_bullet(doc, "Accesos directos a inventario y tareas frecuentes.")

heading(doc, "5. Inventario, lotes y vencimientos", 1)
heading(doc, "5.1 Catálogo e inventario", 2)
body(doc, "Cada medicamento mantiene código interno, código de barras, categoría, laboratorio, presentación, precio de compra, precio de venta, margen, stock actual y stock mínimo por farmacia.")
add_step(doc, "Buscar", "Use la barra única de búsqueda para localizar por nombre, SKU, laboratorio o código.")
add_step(doc, "Revisar", "Observe el estado En stock, Stock crítico o Agotado.")
add_step(doc, "Mantener", "Los administradores pueden crear, editar o retirar medicamentos; cada cambio queda disponible para auditoría.")
heading(doc, "5.2 Lotes y FEFO", 2)
add_bullet(doc, "La recepción de compras registra número de lote, vencimiento, cantidad y costo.")
add_bullet(doc, "Las ventas consumen primero el lote con fecha de vencimiento más próxima (FEFO).")
add_bullet(doc, "Los lotes vencidos quedan bloqueados para ventas y transferencias.")
add_bullet(doc, "Las alertas muestran lotes vencidos o próximos a vencer dentro del periodo configurado.")

heading(doc, "6. Proveedores y compras", 1)
heading(doc, "6.1 Proveedores", 2)
body(doc, "El directorio registra razón social, RUC, contacto, teléfono, correo, ciudad y estado. Sirve como base para las órdenes y devoluciones.")
heading(doc, "6.2 Ciclo de compra", 2)
add_step(doc, "Crear orden", "Seleccione farmacia, proveedor y productos; indique cantidad, costo y observaciones.")
add_step(doc, "Recibir", "Registre por producto la cantidad recibida, lote y vencimiento. Se admiten recepciones parciales.")
add_step(doc, "Actualizar stock", "La recepción crea o incrementa el lote y actualiza el inventario de forma transaccional.")
add_step(doc, "Cerrar orden", "Cuando todo fue recibido, el estado cambia a Recibida; si falta mercancía, permanece Parcial.")
add_step(doc, "Anular o devolver", "Un administrador debe indicar el motivo. La devolución al proveedor descuenta el lote y genera trazabilidad.")

heading(doc, "7. Transferencias entre farmacias", 1)
body(doc, "Las transferencias conservan la trazabilidad del medicamento y lote entre una farmacia de origen y otra de destino.")
add_step(doc, "Solicitar", "Seleccione origen, destino y cantidades disponibles por lote.")
add_step(doc, "Aprobar", "El administrador revisa y autoriza la solicitud.")
add_step(doc, "Despachar", "El sistema descuenta stock y lote del origen en una transacción.")
add_step(doc, "Recibir", "El sistema crea o incrementa el mismo lote en el destino y registra el movimiento.")
callout(doc, "Control", "No se permite transferir a la misma sucursal, superar la cantidad disponible ni utilizar lotes vencidos.", "FFF4D8")

heading(doc, "8. Caja y punto de venta", 1)
heading(doc, "8.1 Caja", 2)
add_step(doc, "Abrir turno", "Seleccione la farmacia e indique el saldo inicial. Cada usuario mantiene su propia sesión de caja.")
add_step(doc, "Registrar movimientos", "Añada ingresos o egresos con monto, forma de pago, referencia y observación.")
add_step(doc, "Cerrar", "Declare el saldo contado. El sistema compara el esperado con el declarado y calcula la diferencia.")
heading(doc, "8.2 Venta", 2)
add_step(doc, "Elegir producto", "Busque el medicamento y verifique stock, precio y sucursal.")
add_step(doc, "Indicar venta", "Registre cantidad, cliente y forma de pago: efectivo, tarjeta o transferencia.")
add_step(doc, "Confirmar", "La operación descuenta existencias por FEFO y guarda usuario, farmacia, lote, fecha y precio.")
add_step(doc, "Controlar efectivo", "Una venta en efectivo requiere una caja abierta para el usuario.")
heading(doc, "8.3 Anulación y devolución del cliente", 2)
add_bullet(doc, "La anulación requiere autorización administrativa y un motivo.")
add_bullet(doc, "La devolución repone la cantidad en el lote original cuando corresponde.")
add_bullet(doc, "Se genera una nota de crédito y queda registrada la persona que autorizó.")

heading(doc, "9. Reportes y Excel", 1)
body(doc, "Los reportes se generan en el servidor desde Supabase, no desde los datos temporales del navegador. Esto reduce diferencias y permite reconstruir el origen de la información.")
add_bullet(doc, "Filtros por rango de fechas, farmacia, usuario y tipo de movimiento.")
add_bullet(doc, "Exportación a Excel mediante botón y confirmación visual del sistema.")
add_bullet(doc, "Datos de ventas, inventario y movimientos según el reporte seleccionado.")
add_bullet(doc, "El archivo incluye encabezados y estructura preparada para análisis y respaldo.")

heading(doc, "10. Auditoría y trazabilidad", 1)
body(doc, "PharmaSys registra eventos relevantes con usuario, farmacia, acción, entidad, identificador, descripción, cambios, fecha e IP cuando está disponible.")
add_table(doc, ["Proceso", "Qué queda registrado"], [
    ("Ventas", "Creación, lotes consumidos, usuario, pago, anulación y devolución."),
    ("Compras", "Orden, recepción parcial o total, lotes, costos, anulación y devolución."),
    ("Inventario", "Entradas, salidas, saldo anterior, saldo nuevo y documento relacionado."),
    ("Transferencias", "Solicitud, aprobación, despacho, recepción, origen y destino."),
    ("Caja", "Apertura, movimientos, cierre, saldo esperado y diferencia."),
    ("Administración", "Cambios en medicamentos, usuarios, permisos y sucursales."),
])

heading(doc, "11. Sucursales y operación multi-farmacia", 1)
add_bullet(doc, "El administrador crea y mantiene farmacias con código, dirección, ciudad, teléfono y responsable.")
add_bullet(doc, "Los usuarios operativos solo reciben información de sus farmacias asignadas.")
add_bullet(doc, "El selector superior cambia el contexto de inventario, ventas, compras y caja.")
add_bullet(doc, "Los administradores pueden consultar y comparar todas las sedes.")

heading(doc, "12. Estados y mensajes del sistema", 1)
body(doc, "Las confirmaciones, motivos, errores y avisos se presentan con componentes visuales de PharmaSys. No se usan cuadros nativos del navegador para acciones como eliminar, reponer, anular o cerrar sesión.")
add_bullet(doc, "Éxito: confirma que la operación fue guardada.")
add_bullet(doc, "Advertencia: solicita revisión antes de una acción sensible.")
add_bullet(doc, "Error: explica por qué la operación no pudo completarse.")
add_bullet(doc, "Confirmación: permite cancelar o continuar sin perder el contexto.")

heading(doc, "13. Flujos diarios recomendados", 1)
heading(doc, "13.1 Apertura", 2)
add_bullet(doc, "Ingresar con Google y comprobar la farmacia seleccionada.")
add_bullet(doc, "Revisar dashboard, alertas de vencimiento y stock crítico.")
add_bullet(doc, "Abrir caja si se realizarán ventas en efectivo.")
heading(doc, "13.2 Operación", 2)
add_bullet(doc, "Recibir compras indicando lote y vencimiento reales.")
add_bullet(doc, "Registrar todas las ventas con su forma de pago correcta.")
add_bullet(doc, "Usar transferencias para mover stock entre farmacias; no realizar ajustes informales.")
heading(doc, "13.3 Cierre", 2)
add_bullet(doc, "Cerrar la caja y revisar diferencias.")
add_bullet(doc, "Exportar reportes necesarios para control interno.")
add_bullet(doc, "Cerrar sesión, especialmente en equipos compartidos.")

heading(doc, "14. Solución de problemas", 1)
add_table(doc, ["Situación", "Qué revisar"], [
    ("Google no permite entrar", "Correo autorizado, cuenta Google seleccionada y configuración OAuth del entorno."),
    ("No aparece un módulo", "Rol asignado. El usuario debe cerrar sesión y volver a entrar después de un cambio."),
    ("No aparece una farmacia", "Asignación activa del usuario a esa sucursal."),
    ("No permite vender en efectivo", "Debe existir una caja abierta para el usuario y farmacia actual."),
    ("No permite vender o transferir", "Stock suficiente, lote vigente y acceso a la sucursal."),
    ("No permite eliminar o anular", "La acción necesita rol administrador o existe una protección de integridad."),
    ("Un reporte no descarga", "Filtros, sesión activa y permiso administrativo cuando aplique."),
])

heading(doc, "15. Alcance técnico actual", 1)
add_bullet(doc, "Backend: Django y API REST autenticada.")
add_bullet(doc, "Base de datos: PostgreSQL administrado en Supabase.")
add_bullet(doc, "Frontend: React con diseño adaptable y diálogos consistentes.")
add_bullet(doc, "Despliegue: servicio web en Render conectado al repositorio GitHub.")
add_bullet(doc, "Autenticación: Google OAuth mediante django-allauth.")
callout(doc, "Facturación electrónica", "La estructura puede ampliarse para cargar certificados y firmar comprobantes. La firma, validación y transmisión oficial al SRI deben completarse y certificarse antes de considerarlas funciones productivas.", "FFF4D8")

heading(doc, "16. Lista de control para administradores", 1)
for text in (
    "Mantener al menos un administrador activo.",
    "Crear usuarios con su correo Google exacto y asignar el rol mínimo necesario.",
    "Verificar sucursales y responsables antes de operar.",
    "Registrar proveedores, productos, precios y mínimos de stock.",
    "Controlar lotes y vencimientos en cada recepción.",
    "Revisar cierres de caja, diferencias y anulaciones.",
    "Exportar reportes periódicos y revisar la auditoría ante inconsistencias.",
):
    add_bullet(doc, text)

doc.core_properties.title = "Manual funcional de PharmaSys"
doc.core_properties.subject = "Guía de operación y funcionalidades"
doc.core_properties.author = "PharmaSys"
doc.core_properties.keywords = "farmacia, inventario, ventas, compras, caja, Supabase"
doc.save(OUT)
print(OUT)
